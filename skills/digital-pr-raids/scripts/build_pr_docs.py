#!/usr/bin/env python3
"""
build_pr_docs.py — Generate the four RAIDS digital-PR deliverables from one JSON config.

Outputs (into the given output dir):
  1. <Client> - Ideation Sheet.docx        (Recon)
  2. <Client> - Dataset.xlsx                (Analyze)   [if dataset_columns provided]
  3. <Client> - Media List.xlsx             (Infiltrate)
  4. <Client> - Pitch & Press Release.docx  (Deploy)
  5. <Client> - Coverage Tracker.xlsx       (Score)

Usage:
  python3 build_pr_docs.py campaign.json /path/to/output_dir

Deps: python-docx, openpyxl  (pip install python-docx openpyxl --break-system-packages)

JSON schema (all keys optional except client; the script renders whatever it's given):
{
  "client": "Kadi Luggage", "url": "https://kadiluggage.com/", "date": "2025-01-06",
  "campaign_type": "Main data-led",            // or "Smaller data-led" / "Reactive"
  "reason": "Initial ideation",
  "campaign_title": "The Hidden Cost of Flying...",
  "headline": "Revealed: The Airlines Charging the Most Hidden Fees in Australia",
  "alt_headlines": ["..."],
  "done_before": "Yes, but not recently with an AU-first focus...",
  "data_collection": ["Analyse baggage/seat/payment fees...", "..."],
  "limitations": ["Fee structures change often", "..."],
  "regional_hooks": ["Domestic vs international", "..."],
  "dream_publications": ["Traveller", "9News", "Finder"],
  "topics_cover": ["..."], "topics_avoid": ["..."],
  "why_care": "Everyone flies and hates surprise fees.",
  "markets": ["Australia (primary)", "UK/US syndication"],
  "outlets": ["Traveller", "Escape", "Finder", "Canstar"],
  "smaller_ideas": ["The Most Expensive Add-Ons in the Sky — ...", "..."],

  "key_findings": ["Malaysia Airlines charged up to 66% more...", "..."],
  "methodology": "We analysed a 6000-8000km route...",
  "quote": {"name": "Harry Saunders", "role": "Co-founder", "text": "This data highlights..."},
  "about_client": "At Kadi, we get it...",
  "credit_line": "https://kadiluggage.com",

  "dataset_columns": ["Ranking", "Airline", "Add-on %"],
  "dataset_rows": [[1, "Malaysia Airlines", "66.18"], ...],

  "media_list": [{"outlet":"Traveller","journalist":"Jane Doe","email":"j@x.com","article_url":"https://...","notes":"covers travel fees"}],

  "subject_lines": ["100,000 flights analysed: the 3-hour slot to avoid delays"],
  "pitch_email": "Hi [First name],\n...",
  "followup_email": "Hi [First name],\n..."
}
"""
import json, os, sys, datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.stderr.write("Install python-docx: pip install python-docx --break-system-packages\n"); sys.exit(2)
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
except ImportError:
    sys.stderr.write("Install openpyxl: pip install openpyxl --break-system-packages\n"); sys.exit(2)

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLUE = RGBColor(0x25, 0x63, 0xEB)
HEADER_FILL = PatternFill("solid", fgColor="1B2A4A")
WHITE = Font(color="FFFFFF", bold=True)


def H(doc, text, size=16, color=NAVY, before=10, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = "Arial"
    return p


def field(doc, label, value):
    p = doc.add_paragraph()
    r = p.add_run(label + ": "); r.bold = True; r.font.name = "Arial"; r.font.size = Pt(11)
    r2 = p.add_run(value if value else "—"); r2.font.name = "Arial"; r2.font.size = Pt(11)


def bullets(doc, items):
    for it in (items or []):
        p = doc.add_paragraph(style="List Bullet"); r = p.add_run(str(it)); r.font.name = "Arial"; r.font.size = Pt(11)


def para(doc, text, bold=False, size=11):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.font.name = "Arial"; r.font.size = Pt(size); return p


# ---------------- 1. IDEATION SHEET ----------------
def build_ideation(c, out):
    doc = Document()
    H(doc, "Ideation Sheet — %s" % c.get("client", ""), 20)
    field(doc, "Client", c.get("client"))
    field(doc, "URL", c.get("url"))
    field(doc, "Date", c.get("date"))
    field(doc, "Reason for ideation", c.get("reason"))
    field(doc, "Campaign type", c.get("campaign_type"))

    H(doc, "Main Data-Led Campaign Idea", 15, BLUE)
    field(doc, "Title", c.get("campaign_title"))
    field(doc, "Headline", c.get("headline"))
    if c.get("alt_headlines"):
        para(doc, "Alternate headlines:", bold=True); bullets(doc, c["alt_headlines"])
    field(doc, "Has it been done before?", c.get("done_before"))
    para(doc, "How we'll collect the data:", bold=True); bullets(doc, c.get("data_collection"))
    para(doc, "Limitations:", bold=True); bullets(doc, c.get("limitations"))
    para(doc, "Regional / additional hooks:", bold=True); bullets(doc, c.get("regional_hooks"))
    para(doc, "Dream publications:", bold=True); bullets(doc, c.get("dream_publications"))
    para(doc, "Outlets that could run it:", bold=True); bullets(doc, c.get("outlets"))
    field(doc, "Topics to cover", ", ".join(c.get("topics_cover", [])) or None)
    field(doc, "Topics to avoid", ", ".join(c.get("topics_avoid", [])) or None)
    field(doc, "Why would anyone care?", c.get("why_care"))
    para(doc, "Markets:", bold=True); bullets(doc, c.get("markets"))

    H(doc, "Other Smaller or Reactive Ideas", 15, BLUE)
    bullets(doc, c.get("smaller_ideas"))

    para(doc, "")
    para(doc, "Next step: post these ideas in Slack and run a quick vote to pick the best one.", bold=True)
    doc.save(out); print("Wrote", out)


# ---------------- 2. DATASET ----------------
def build_dataset(c, out):
    cols = c.get("dataset_columns")
    if not cols:
        return
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Dataset"
    ws.append(cols)
    for ci in range(1, len(cols) + 1):
        cell = ws.cell(1, ci); cell.fill = HEADER_FILL; cell.font = WHITE
    for row in c.get("dataset_rows", []):
        ws.append(row)
    for ci, col in enumerate(cols, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = max(14, min(40, len(str(col)) + 6))
    wb.save(out); print("Wrote", out)


# ---------------- 3. MEDIA LIST ----------------
def build_media_list(c, out):
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Media List"
    cols = ["Outlet", "Journalist Name", "Email Address", "Article URL", "Notes"]
    ws.append(cols)
    for ci in range(1, len(cols) + 1):
        cell = ws.cell(1, ci); cell.fill = HEADER_FILL; cell.font = WHITE
    for m in c.get("media_list", []):
        ws.append([m.get("outlet", ""), m.get("journalist", ""), m.get("email", ""),
                   m.get("article_url", ""), m.get("notes", "")])
    # pad to a usable working sheet
    for _ in range(max(0, 25 - len(c.get("media_list", [])))):
        ws.append(["", "", "", "", ""])
    widths = [26, 22, 30, 46, 40]
    for ci, w in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = w
    wb.save(out); print("Wrote", out)


# ---------------- 4. PITCH & PRESS RELEASE ----------------
def build_pitch(c, out):
    doc = Document()
    H(doc, "%s — Pitch & Press Release" % c.get("client", ""), 18)

    H(doc, "Subject line options", 14, BLUE)
    bullets(doc, c.get("subject_lines"))

    H(doc, "Short pitch email", 14, BLUE)
    for line in (c.get("pitch_email") or "").split("\n"):
        para(doc, line)

    H(doc, "Headline options", 14, BLUE)
    bullets(doc, [c.get("headline")] + c.get("alt_headlines", []) if c.get("headline") else c.get("alt_headlines"))

    H(doc, "Press Release", 16)
    if c.get("headline"):
        para(doc, c["headline"], bold=True, size=14)
    para(doc, "Key findings:", bold=True); bullets(doc, c.get("key_findings"))

    # dataset table inline
    if c.get("dataset_columns") and c.get("dataset_rows"):
        cols = c["dataset_columns"]; rows = c["dataset_rows"]
        t = doc.add_table(rows=1, cols=len(cols)); t.style = "Light Grid Accent 1"
        for ci, col in enumerate(cols):
            t.rows[0].cells[ci].text = str(col)
        for r in rows:
            cells = t.add_row().cells
            for ci, v in enumerate(r):
                cells[ci].text = str(v)

    if c.get("quote"):
        q = c["quote"]
        para(doc, "")
        para(doc, "%s, %s, commented:" % (q.get("name", ""), q.get("role", "")), bold=True)
        para(doc, "“%s”" % q.get("text", ""))

    para(doc, ""); para(doc, "— END —", bold=True)

    H(doc, "Data & Methodology", 14, BLUE)
    para(doc, c.get("methodology") or "—")

    H(doc, "About %s" % c.get("client", ""), 14, BLUE)
    para(doc, c.get("about_client") or "—")

    H(doc, "Editor's Notes", 14, BLUE)
    para(doc, "If you use this research, please add a linked credit to: %s" %
         (c.get("credit_line") or c.get("url") or ""))

    if c.get("followup_email"):
        H(doc, "Follow-up email (send after 5–7 days)", 14, BLUE)
        for line in c["followup_email"].split("\n"):
            para(doc, line)

    doc.save(out); print("Wrote", out)


# ---------------- 5. COVERAGE TRACKER ----------------
def build_tracker(c, out):
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Coverage"
    cols = ["Type Of Campaign", "Campaign Name", "Date Published", "Link Type",
            "Domain Rating (DR)", "Linked to (Homepage, category)", "Client Link",
            "Link to Article", "Anchor Text"]
    ws.append(cols)
    for ci in range(1, len(cols) + 1):
        cell = ws.cell(1, ci); cell.fill = HEADER_FILL; cell.font = WHITE
    # seed one example row with the campaign so the attendee sees the format
    ws.append([c.get("campaign_type", ""), c.get("campaign_title", c.get("headline", "")),
               "", "Clean", "", "Homepage", c.get("url", ""), "", c.get("client", "")])
    for _ in range(30):
        ws.append([""] * len(cols))
    widths = [18, 22, 16, 12, 14, 26, 26, 50, 22]
    for ci, w in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = w
    wb.save(out); print("Wrote", out)


def main():
    if len(sys.argv) < 3:
        sys.stderr.write("Usage: build_pr_docs.py campaign.json output_dir\n"); sys.exit(1)
    cfg = json.load(open(sys.argv[1], encoding="utf-8"))
    outdir = sys.argv[2]; os.makedirs(outdir, exist_ok=True)
    client = cfg.get("client", "Client").replace("/", "-")
    build_ideation(cfg, os.path.join(outdir, "%s - Ideation Sheet.docx" % client))
    build_dataset(cfg, os.path.join(outdir, "%s - Dataset.xlsx" % client))
    build_media_list(cfg, os.path.join(outdir, "%s - Media List.xlsx" % client))
    build_pitch(cfg, os.path.join(outdir, "%s - Pitch & Press Release.docx" % client))
    build_tracker(cfg, os.path.join(outdir, "%s - Coverage Tracker.xlsx" % client))
    print("\nAll deliverables written to:", outdir)


if __name__ == "__main__":
    main()
