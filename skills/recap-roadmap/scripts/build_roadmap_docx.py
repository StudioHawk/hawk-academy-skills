#!/usr/bin/env python3
"""
build_roadmap_docx.py — Render the 6-Month Roadmap deliverable from a JSON definition.

Usage:
    python3 build_roadmap_docx.py roadmap.json /path/to/output_dir/

Deps: python-docx  (pip install python-docx --break-system-packages)

JSON schema (only "business" + "north_star" required):
{
  "business": "Fit For Life Health & Fitness",
  "domain": "fitforlifefitness.com.au",
  "date": "2026-06-12",
  "north_star": "Double organic enquiries within 6 months",
  "baseline": "Where you are now — 2-3 sentences from the audit.",
  "first_two_weeks": ["Fix meta descriptions on money pages", "..."],
  "phases": [
    {
      "name": "Phase 1 — Foundations & Quick Wins (Months 0–2)",
      "objective": "Reclaim equity and ship fast high-impact fixes.",
      "actions": ["Refresh /personal-training (meta, H1, FAQ schema)", "..."],
      "owner": "Owner / team / agency",
      "kpis": ["+20% impressions on money pages", "..."],
      "dependencies": ["GA4 + Clarity live"]
    }
  ],
  "monthly_focus": {"1": "Reclaim equity", "2": "...", "6": "..."}
}
"""
import json, sys, os
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("Install python-docx: pip install python-docx --break-system-packages\n"); sys.exit(2)

NAVY=RGBColor(0x1B,0x2A,0x4A); BLUE=RGBColor(0x25,0x63,0xEB); GREY=RGBColor(0x64,0x74,0x8B); DARK=RGBColor(0x1E,0x29,0x3B)
PHASE_FILL=["DBEAFE","FEF3C7","DCFCE7"]  # phase 1/2/3 tints

def run(p,t,b=False,sz=11,c=DARK,i=False):
    r=p.add_run(t); r.bold=b; r.italic=i; r.font.name="Arial"; r.font.size=Pt(sz); r.font.color.rgb=c; return r
def H(doc,t,sz=16,c=NAVY,before=12,after=4):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    run(p,t,True,sz,c); return p
def para(doc,t,b=False,sz=11,c=DARK,i=False):
    p=doc.add_paragraph(); run(p,t,b,sz,c,i); return p
def bullets(doc,items):
    for it in (items or []):
        p=doc.add_paragraph(style="List Bullet"); run(p,str(it))
def shade(cell,hexfill):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),hexfill); tcPr.append(sh)

def main():
    if len(sys.argv)<3: sys.stderr.write("Usage: build_roadmap_docx.py roadmap.json output_dir\n"); sys.exit(1)
    c=json.load(open(sys.argv[1],encoding="utf-8")); outdir=sys.argv[2]; os.makedirs(outdir,exist_ok=True)
    doc=Document(); doc.styles["Normal"].font.name="Arial"; doc.styles["Normal"].font.size=Pt(11)
    H(doc,"6-Month SEO Roadmap",26,NAVY,before=0)
    p=doc.add_paragraph(); run(p,c.get("business",""),True,14,BLUE)
    meta="  ·  ".join([x for x in [c.get("domain",""),c.get("date","")] if x])
    if meta: p=doc.add_paragraph(); run(p,meta,sz=10,c=GREY)
    # north star
    H(doc,"North-star goal",14,BLUE)
    para(doc,c.get("north_star","—"),b=True)
    if c.get("baseline"):
        H(doc,"Where you are now",14,BLUE); para(doc,c["baseline"])
    if c.get("first_two_weeks"):
        H(doc,"First 2 weeks — kickstart",14,BLUE); bullets(doc,c["first_two_weeks"])
    # phases
    for idx,ph in enumerate(c.get("phases",[])):
        H(doc,ph.get("name","Phase"),15,BLUE)
        if ph.get("objective"): para(doc,"Objective: "+ph["objective"],b=True)
        # actions table
        if ph.get("actions"):
            para(doc,"Actions:",b=True); bullets(doc,ph["actions"])
        cols=[]
        if ph.get("owner"): cols.append(("Owner",ph["owner"]))
        if ph.get("kpis"): cols.append(("KPIs", "; ".join(ph["kpis"]) if isinstance(ph["kpis"],list) else ph["kpis"]))
        if ph.get("dependencies"): cols.append(("Dependencies","; ".join(ph["dependencies"]) if isinstance(ph["dependencies"],list) else ph["dependencies"]))
        if cols:
            t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
            fill=PHASE_FILL[idx % len(PHASE_FILL)]
            for k,v in cols:
                r=t.add_row().cells; shade(r[0],fill)
                run(r[0].paragraphs[0],k,True,10); run(r[1].paragraphs[0],str(v),sz=10)
    # monthly focus
    mf=c.get("monthly_focus")
    if mf:
        H(doc,"Protect time for — the one thing each month",14,BLUE)
        for m in sorted(mf, key=lambda x:int(x)):
            para(doc,f"Month {m}: {mf[m]}")
    out=os.path.join(outdir,"6-Month-Roadmap-%s.docx"%c.get("business","client").replace("/","-").replace(" ","-"))
    doc.save(out); print("Wrote",out)

if __name__=="__main__":
    main()
