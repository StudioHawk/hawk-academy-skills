#!/usr/bin/env python3
"""
build_ecommerce_docx.py

Builds a branded StudioHawk eCommerce SEO audit Word document from a JSON
definition written by the skill.

Usage:
    python3 build_ecommerce_docx.py audit.json "eCommerce SEO Audit - Store - 2026-08-02.docx"

EVERY SECTION IS OPTIONAL except the handful marked with *. Include what the
crawl actually supports and omit what it does not; the document renders
cleanly either way. Do not invent rows to fill a section out.

JSON schema
-----------

{
  "store_name": "Example Store",                    *
  "store_url": "https://store.example.com",         *
  "locale": "AU",                                      // optional, defaults to "AU"
  "date": "2026-08-02",                                // optional, defaults to today
  "platform": "Shopify",                               // optional

  "crawl_summary": {                                   // optional
    "pages_crawled": 240,
    "categories": 18,
    "products": 96,
    "blogs": 22,
    "max_depth": 3,
    "limitations": [
      "Crawl capped at depth 3, so deep sub-categories were not reached."
    ]
  },

  "executive_summary": {                            *
    "headline": "Strong catalogue, weak schema and thin category copy.",
    "schema_headline": "12 of 40 product pages are Merchant-ready",
    "biggest_category_gap": "8 linen-dress SKUs with no /collections/linen-dresses",
    "robots_verdict": "No money pages blocked",
    "top_actions": [                                   // up to 3, P0 only
      "Add brand to Product schema across all 40 PDPs",
      "Build a linen dresses category page",
      "Publish a returns policy page"
    ]
  },

  "scorecard": [                                    *  // one row per audit area
    {"area": "Category coverage", "status": "Partial", "note": "3 clear gaps"}
    // status: Pass / Partial / Fail / Not assessable
  ],

  "category_coverage": {                               // optional
    "commentary": "...",
    "existing": ["dresses", "tops", "knitwear"],
    "gaps": [
      {"theme": "Linen dresses",
       "evidence": "8 linen-dress SKUs, no matching category page",
       "suggested_url": "/collections/linen-dresses",
       "priority": "P1"}
    ]
  },

  "category_copy": {                                   // optional
    "commentary": "...",
    "rows": [
      {"url": "/collections/dresses", "abf": "No", "abf_words": 0,
       "btf": "Yes", "btf_words": 210,
       "note": "Add a 60-word keyword-led intro above the grid"}
    ]
  },

  "product_schema": {                                  // optional
    "commentary": "...",
    "product_pages": 40, "with_schema": 32,
    "merchant_ready": 12, "no_schema": 8,
    "rows": [
      {"url": "/products/linen-midi",
       "has_schema": true, "merchant_ready": false,
       "missing_required": ["brand"], "missing_recommended": ["review"]}
    ]
  },

  "product_pages": {                                   // optional
    "commentary": "...",
    "rows": [
      {"url": "/products/linen-midi", "h1": "Yes", "intro": "No",
       "specs": "Yes", "scannable": "Yes",
       "fix": "Add a 40-word intro above the spec block"}
    ]
  },

  "blog": {                                            // optional
    "commentary": "...",
    "rows": [
      {"url": "/blogs/news/how-to-style-linen", "tie_back": "Yes",
       "early_tie_back": "No",
       "fix": "Move the category link into the opening two paragraphs"}
    ]
  },

  "key_pages": {                                       // optional
    "commentary": "...",
    "rows": [
      {"page": "Returns / refunds", "status": "Missing", "url": "",
       "note": "No returns policy page found in the crawl"}
      // status: Found / Missing / Not assessable
    ]
  },

  "robots": {                                          // optional
    "verdict": "No money pages blocked",
    "commentary": "...",
    "rows": [
      {"path": "/search", "user_agent": "*", "severity": "benign",
       "impact": "Internal search, correctly excluded"}
      // severity: critical / high / review / benign
    ]
  },

  "action_plan": [                                  *
    {"priority": "P0", "action": "Publish a returns policy page",
     "area": "Trust pages",
     "impact": "Removes a buyer-trust and eligibility blocker"}
  ]
}
"""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx is not installed. Install with:\n"
        "    pip install python-docx --break-system-packages\n"
    )
    sys.exit(2)


# --- StudioHawk palette ------------------------------------------------------
# Picton Blue + black is the current StudioHawk brand. Blue is used for rules
# and table-header fills with BLACK text on top, because white on light blue
# fails contrast. Heading text stays black for the same reason.

HAWK_BLUE_HEX = "3DB7E9"
HAWK_BLUE = RGBColor(0x3D, 0xB7, 0xE9)
HAWK_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
HAWK_GREY = RGBColor(0x55, 0x55, 0x55)
HAWK_RED = RGBColor(0xC0, 0x2B, 0x2B)
HAWK_GREEN = RGBColor(0x1F, 0x7A, 0x3E)
HAWK_AMBER = RGBColor(0xB0, 0x72, 0x1F)

BODY_FONT = "Archivo"          # official StudioHawk body face
FALLBACK_FONT = "Calibri"      # if Archivo is not installed, Word substitutes


# --- Generic helpers ---------------------------------------------------------

def _shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _rule(doc, colour_hex: str = HAWK_BLUE_HEX, size: int = 12) -> None:
    """A thin brand-coloured horizontal rule."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour_hex)
    borders.append(bottom)
    p_pr.append(borders)


def _heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = HAWK_BLACK
        run.font.name = BODY_FONT
        if level == 1:
            run.font.size = Pt(19)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)


def _para(doc, text: str, *, bold: bool = False, italic: bool = False,
          color=None, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = BODY_FONT
    if color is not None:
        run.font.color.rgb = color


def _label_value(doc, label: str, value: str) -> None:
    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    label_run.font.name = BODY_FONT
    value_run = p.add_run(str(value))
    value_run.font.size = Pt(11)
    value_run.font.name = BODY_FONT


def _set_header_cell(cell, text: str) -> None:
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = BODY_FONT
            run.font.color.rgb = HAWK_BLACK      # black on blue, not white
    _shade_cell(cell, HAWK_BLUE_HEX)


def _set_cell_text(cell, text: str, *, bold: bool = False, color=None, size: int = 10) -> None:
    cell.text = str(text)
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = BODY_FONT
            if color is not None:
                run.font.color.rgb = color


def _status_colour(status: str) -> RGBColor:
    s = (status or "").strip().lower()
    if s in ("pass", "yes", "found", "strong", "benign", "ok"):
        return HAWK_GREEN
    if s in ("partial", "review", "amber", "some"):
        return HAWK_AMBER
    if s in ("fail", "no", "missing", "critical", "high"):
        return HAWK_RED
    return HAWK_GREY if s in ("not assessable", "unknown", "n/a") else HAWK_BLACK


def _priority_fill(priority: str) -> str:
    p = (priority or "").upper().strip()
    return {"P0": "F4CFCF", "P1": "F8E4C3", "P2": "D7E9D4"}.get(p, "EEEEEE")


def _new_table(doc, headers: list[str]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, label in enumerate(headers):
        _set_header_cell(table.rows[0].cells[i], label)
    return table


def _yn(value) -> str:
    """Normalise booleans and strings into a Yes/No/blank cell value."""
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if value is None:
        return ""
    return str(value)


class _SectionNumber:
    """Numbers sections as they render, so omitted areas leave no gap."""

    def __init__(self) -> None:
        self.n = 0

    def __call__(self) -> int:
        self.n += 1
        return self.n


# --- Sections ----------------------------------------------------------------

def _add_cover(doc, data: dict) -> None:
    title = doc.add_paragraph()
    run = title.add_run("eCommerce SEO Audit")
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = BODY_FONT
    run.font.color.rgb = HAWK_BLACK

    sub = doc.add_paragraph()
    sub_run = sub.add_run(data.get("store_name", "Online Store"))
    sub_run.bold = True
    sub_run.font.size = Pt(18)
    sub_run.font.name = BODY_FONT
    sub_run.font.color.rgb = HAWK_BLACK

    _rule(doc)

    bits = [data.get("store_url", ""), data.get("locale", "AU")]
    if data.get("platform"):
        bits.append(data["platform"])
    bits.append(data.get("date", date.today().isoformat()))

    meta = doc.add_paragraph()
    meta_run = meta.add_run("  |  ".join(b for b in bits if b))
    meta_run.font.size = Pt(10)
    meta_run.font.name = BODY_FONT
    meta_run.font.color.rgb = HAWK_GREY
    doc.add_paragraph()


def _add_crawl_summary(doc, data: dict) -> None:
    cs = data.get("crawl_summary")
    if not cs:
        return
    _heading(doc, "Crawl Coverage", level=2)
    parts = []
    for key, label in (("pages_crawled", "pages"), ("categories", "category pages"),
                       ("products", "product pages"), ("blogs", "blog posts")):
        if cs.get(key) is not None:
            parts.append(f"{cs[key]} {label}")
    if parts:
        _label_value(doc, "Crawled", ", ".join(parts))
    if cs.get("max_depth") is not None:
        _label_value(doc, "Max depth", cs["max_depth"])

    limitations = cs.get("limitations") or []
    if limitations:
        _para(doc, "Limitations", bold=True)
        for item in limitations:
            doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()


def _add_exec_summary(doc, data: dict) -> None:
    es = data.get("executive_summary") or {}
    _heading(doc, "Executive Summary", level=1)

    if es.get("headline"):
        _para(doc, es["headline"], bold=True)
        doc.add_paragraph()

    for key, label in (
        ("schema_headline", "Product schema"),
        ("biggest_category_gap", "Biggest category gap"),
        ("robots_verdict", "robots.txt"),
    ):
        if es.get(key):
            _label_value(doc, label, es[key])

    actions = es.get("top_actions") or []
    if actions:
        doc.add_paragraph()
        _para(doc, "Top P0 actions", bold=True)
        for a in actions[:3]:
            doc.add_paragraph(a, style="List Bullet")
    doc.add_paragraph()


def _add_scorecard(doc, data: dict, num) -> None:
    rows = data.get("scorecard") or []
    _heading(doc, f"{num()}. Audit Scorecard", level=1)
    if not rows:
        _para(doc, "No scorecard provided.", italic=True, color=HAWK_GREY)
        return

    table = _new_table(doc, ["Audit area", "Status", "Note"])
    for row in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], row.get("area", ""), bold=True)
        status = row.get("status", "")
        _set_cell_text(cells[1], status, bold=True, color=_status_colour(status))
        _set_cell_text(cells[2], row.get("note", "") or "")
    doc.add_paragraph()


def _add_category_coverage(doc, data: dict, num) -> None:
    section = data.get("category_coverage")
    if not section:
        return
    _heading(doc, f"{num()}. Category Page Coverage", level=1)
    if section.get("commentary"):
        _para(doc, section["commentary"])

    existing = section.get("existing") or []
    if existing:
        _label_value(doc, "Category pages found", ", ".join(str(e) for e in existing))

    gaps = section.get("gaps") or []
    if not gaps:
        _para(doc, "No category gaps identified.", italic=True, color=HAWK_GREY)
        doc.add_paragraph()
        return

    doc.add_paragraph()
    table = _new_table(doc, ["Missing category", "Evidence", "Suggested URL", "Priority"])
    for gap in gaps:
        cells = table.add_row().cells
        _set_cell_text(cells[0], gap.get("theme", ""), bold=True)
        _set_cell_text(cells[1], gap.get("evidence", ""))
        _set_cell_text(cells[2], gap.get("suggested_url", ""))
        priority = (gap.get("priority") or "").upper()
        _set_cell_text(cells[3], priority, bold=True)
        if priority:
            _shade_cell(cells[3], _priority_fill(priority))
    doc.add_paragraph()


def _add_category_copy(doc, data: dict, num) -> None:
    section = data.get("category_copy")
    if not section:
        return
    _heading(doc, f"{num()}. Category Page Copy (above and below the fold)", level=1)
    if section.get("commentary"):
        _para(doc, section["commentary"])
    _para(doc,
          "ABF and BTF word counts are heuristics taken from the crawl. Treat a clear "
          "zero as a strong flag and spot-check anything borderline.",
          italic=True, color=HAWK_GREY, size=10)

    rows = section.get("rows") or []
    if not rows:
        _para(doc, "No category pages assessed.", italic=True, color=HAWK_GREY)
        doc.add_paragraph()
        return

    doc.add_paragraph()
    table = _new_table(doc, ["Category page", "ABF", "Words", "BTF", "Words", "Recommendation"])
    for row in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], row.get("url", ""))
        abf = _yn(row.get("abf"))
        _set_cell_text(cells[1], abf, bold=True, color=_status_colour(abf))
        _set_cell_text(cells[2], row.get("abf_words", ""))
        btf = _yn(row.get("btf"))
        _set_cell_text(cells[3], btf, bold=True, color=_status_colour(btf))
        _set_cell_text(cells[4], row.get("btf_words", ""))
        _set_cell_text(cells[5], row.get("note", "") or "")
    doc.add_paragraph()


def _add_product_schema(doc, data: dict, num) -> None:
    section = data.get("product_schema")
    if not section:
        return
    _heading(doc, f"{num()}. Product Schema (Google Merchant listings)", level=1)
    if section.get("commentary"):
        _para(doc, section["commentary"])

    for key, label in (("product_pages", "Product pages assessed"),
                       ("with_schema", "With Product schema"),
                       ("merchant_ready", "Merchant-ready (all required fields)"),
                       ("no_schema", "No Product schema at all")):
        if section.get(key) is not None:
            _label_value(doc, label, section[key])

    _para(doc,
          "Required: name, image, description, sku/gtin/mpn, brand, offers.price, "
          "offers.priceCurrency, offers.availability. Missing a required field blocks "
          "Merchant listing eligibility. Recommended: aggregateRating, review. Missing "
          "those forgoes review rich results but does not block eligibility.",
          italic=True, color=HAWK_GREY, size=10)

    rows = section.get("rows") or []
    if not rows:
        doc.add_paragraph()
        return

    doc.add_paragraph()
    table = _new_table(doc, ["Product page", "Schema", "Merchant-ready", "Missing required", "Missing recommended"])
    for row in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], row.get("url", ""))
        has_schema = _yn(row.get("has_schema"))
        _set_cell_text(cells[1], has_schema, bold=True, color=_status_colour(has_schema))
        ready = _yn(row.get("merchant_ready"))
        _set_cell_text(cells[2], ready, bold=True, color=_status_colour(ready))
        _set_cell_text(cells[3], ", ".join(row.get("missing_required") or []) or "None")
        _set_cell_text(cells[4], ", ".join(row.get("missing_recommended") or []) or "None")
    doc.add_paragraph()


def _add_product_pages(doc, data: dict, num) -> None:
    section = data.get("product_pages")
    if not section:
        return
    _heading(doc, f"{num()}. Product Page Best Practice", level=1)
    if section.get("commentary"):
        _para(doc, section["commentary"])

    rows = section.get("rows") or []
    if not rows:
        _para(doc, "No product pages assessed.", italic=True, color=HAWK_GREY)
        doc.add_paragraph()
        return

    doc.add_paragraph()
    table = _new_table(doc, ["Product page", "H1", "Intro", "Specs", "Scannable", "Fix"])
    for row in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], row.get("url", ""))
        for i, key in enumerate(("h1", "intro", "specs", "scannable"), start=1):
            value = _yn(row.get(key))
            _set_cell_text(cells[i], value, bold=True, color=_status_colour(value))
        _set_cell_text(cells[5], row.get("fix", "") or "")
    doc.add_paragraph()


def _add_blog(doc, data: dict, num) -> None:
    section = data.get("blog")
    if not section:
        return
    _heading(doc, f"{num()}. Blog Content Tie-back", level=1)
    if section.get("commentary"):
        _para(doc, section["commentary"])
    _para(doc,
          "The rule: every blog post should link to a category or product page, "
          "ideally in the opening lines.",
          italic=True, color=HAWK_GREY, size=10)

    rows = section.get("rows") or []
    if not rows:
        _para(doc, "No blog posts assessed.", italic=True, color=HAWK_GREY)
        doc.add_paragraph()
        return

    doc.add_paragraph()
    table = _new_table(doc, ["Blog post", "Ties back", "Early tie-back", "Fix"])
    for row in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], row.get("url", ""))
        tie = _yn(row.get("tie_back"))
        _set_cell_text(cells[1], tie, bold=True, color=_status_colour(tie))
        early = _yn(row.get("early_tie_back"))
        _set_cell_text(cells[2], early, bold=True, color=_status_colour(early))
        _set_cell_text(cells[3], row.get("fix", "") or "")
    doc.add_paragraph()


def _add_key_pages(doc, data: dict, num) -> None:
    section = data.get("key_pages")
    if not section:
        return
    _heading(doc, f"{num()}. Trust and Seasonal Pages", level=1)
    if section.get("commentary"):
        _para(doc, section["commentary"])

    rows = section.get("rows") or []
    if not rows:
        _para(doc, "No key pages assessed.", italic=True, color=HAWK_GREY)
        doc.add_paragraph()
        return

    doc.add_paragraph()
    table = _new_table(doc, ["Page", "Status", "URL", "Note"])
    for row in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], row.get("page", ""), bold=True)
        status = row.get("status", "")
        _set_cell_text(cells[1], status, bold=True, color=_status_colour(status))
        _set_cell_text(cells[2], row.get("url", "") or "")
        _set_cell_text(cells[3], row.get("note", "") or "")
    doc.add_paragraph()


def _add_robots(doc, data: dict, num) -> None:
    section = data.get("robots")
    if not section:
        return
    _heading(doc, f"{num()}. robots.txt", level=1)
    if section.get("verdict"):
        _label_value(doc, "Verdict", section["verdict"])
    if section.get("commentary"):
        _para(doc, section["commentary"])

    rows = section.get("rows") or []
    if not rows:
        doc.add_paragraph()
        return

    doc.add_paragraph()
    table = _new_table(doc, ["Disallow", "User-agent", "Severity", "Impact"])
    severity_order = {"critical": 0, "high": 1, "review": 2, "benign": 3}
    rows_sorted = sorted(rows, key=lambda r: severity_order.get((r.get("severity") or "").lower(), 4))
    for row in rows_sorted:
        cells = table.add_row().cells
        _set_cell_text(cells[0], row.get("path", ""))
        _set_cell_text(cells[1], row.get("user_agent", "*"))
        severity = row.get("severity", "")
        _set_cell_text(cells[2], severity, bold=True, color=_status_colour(severity))
        _set_cell_text(cells[3], row.get("impact", "") or "")
    doc.add_paragraph()


def _add_action_plan(doc, data: dict, num) -> None:
    actions = data.get("action_plan") or []
    _heading(doc, f"{num()}. Prioritised Action Plan", level=1)
    _para(doc,
          "P0 = fix now, P1 = fix this month, P2 = fix this quarter. "
          "Ordered by impact, not by topic.",
          color=HAWK_GREY)

    if not actions:
        _para(doc, "No actions provided.", italic=True, color=HAWK_GREY)
        return

    order = {"P0": 0, "P1": 1, "P2": 2}
    actions_sorted = sorted(actions, key=lambda a: order.get((a.get("priority") or "").upper(), 99))

    doc.add_paragraph()
    table = _new_table(doc, ["Priority", "Action", "Area", "Expected impact"])
    for a in actions_sorted:
        cells = table.add_row().cells
        priority = (a.get("priority") or "").upper()
        _set_cell_text(cells[0], priority, bold=True)
        _shade_cell(cells[0], _priority_fill(priority))
        _set_cell_text(cells[1], a.get("action", ""))
        _set_cell_text(cells[2], a.get("area", "") or "")
        _set_cell_text(cells[3], a.get("impact", "") or "")
    doc.add_paragraph()


def _add_footer(doc) -> None:
    _rule(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = footer.add_run(
        "Generated by the Hawk Academy eCommerce Auditor. Findings are drawn from a "
        "point-in-time crawl. Copy-length and intro checks are heuristics, so spot-check "
        "borderline calls against the live site before sending this to a client."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = BODY_FONT
    run.font.color.rgb = HAWK_GREY


# --- Main --------------------------------------------------------------------

def build_doc(data: dict, output_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    # Word falls back to Calibri automatically when Archivo is not installed.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:cs"), FALLBACK_FONT)

    num = _SectionNumber()
    _add_cover(doc, data)
    _add_crawl_summary(doc, data)
    _add_exec_summary(doc, data)
    _add_scorecard(doc, data, num)
    _add_category_coverage(doc, data, num)
    _add_category_copy(doc, data, num)
    _add_product_schema(doc, data, num)
    _add_product_pages(doc, data, num)
    _add_blog(doc, data, num)
    _add_key_pages(doc, data, num)
    _add_robots(doc, data, num)
    _add_action_plan(doc, data, num)
    _add_footer(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> int:
    if len(sys.argv) != 3:
        sys.stderr.write(
            "Usage: python3 build_ecommerce_docx.py audit.json output.docx\n"
        )
        return 1

    audit_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not audit_path.exists():
        sys.stderr.write(f"ERROR: input file not found: {audit_path}\n")
        return 1
    try:
        data = json.loads(audit_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        sys.stderr.write(f"ERROR: invalid JSON in {audit_path}: {e}\n")
        return 1

    for required in ("store_name", "store_url"):
        if not data.get(required):
            sys.stderr.write(f"ERROR: audit.json is missing required key: {required}\n")
            return 1

    build_doc(data, output_path)
    print(f"Wrote {output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
