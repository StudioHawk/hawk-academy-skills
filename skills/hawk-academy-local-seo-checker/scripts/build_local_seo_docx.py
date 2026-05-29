#!/usr/bin/env python3
"""
build_local_seo_docx.py

Builds a Hawk Academy Local SEO audit Word document from a JSON definition.

Usage:
    python3 build_local_seo_docx.py audit.json output.docx

JSON schema (keys marked * are required):

{
  "business_name": "Smith & Co Plumbing",         *
  "locality": "Melbourne CBD",                    *
  "locale": "AU",                                    // optional, defaults to "AU"
  "date": "2026-05-18",                              // optional, defaults to today
  "executive_summary": {                          *
    "gbp_score_points": 18,                       *  // out of 25
    "gbp_score_pct": 72,                          *  // 0-100
    "gbp_bucket": "Good with gaps",               *  // Strong / Good with gaps / Material gaps / Failing
    "nap_summary": "3 inconsistencies across 10 directories",  *
    "review_headline": "4.6 from 87 reviews",     *
    "three_pack_rate": "2 of 5 target keywords",  *
    "top_actions": [                              *  // up to 3 lines, P0-only
      "Update Yelp AU phone to +61 3 9876 5432",
      "Claim Apple Maps listing",
      "Respond to last 14 reviews"
    ]
  },
  "gbp_scorecard": [                              *  // 13 rows expected
    {
      "field": "Business name",                   *
      "weight": 3,                                *
      "status": "Pass",                           *  // Pass / Partial / Fail / Not provided
      "note": "Matches canonical name"               // optional one-liner
    }
  ],
  "nap_report": [                                 *
    {
      "directory": "Yelp Australia",              *
      "status": "Inconsistent",                   *  // Match / Inconsistent / Missing
      "name_on_listing": "Smith Co Plumbers",
      "address_on_listing": "12 Bourke St, Melbourne VIC 3000",
      "phone_on_listing": "(03) 9123 4567",
      "discrepancy": "Phone mismatch",
      "fix_instruction": "Update Yelp AU phone to +61 3 9876 5432 to match canonical listing."
    }
  ],
  "review_analysis": {                            *
    "total_reviews": 87,                          *
    "avg_rating": 4.6,                            *
    "reviews_last_90d": 12,                          // optional
    "velocity_per_month": 4.0,                       // optional, derived
    "velocity_grade": "Healthy",                     // Strong / Healthy / Soft / Stalled
    "response_rate_pct": 35,                         // optional
    "response_grade": "Poor",                        // Strong / Acceptable / Poor
    "sentiment_note": "Recent reviews positive, two flagged service delays",
    "commentary": "Overall a healthy review profile...",  *
    "competitors": [                                  // optional; omit if no data
      {"name": "Rival Co", "domain": "rival.com.au", "reviews": 142, "avg_rating": 4.7}
    ]
  },
  "rankings": {                                   *
    "three_pack_rate_pct": 40,                    *  // 0-100
    "data_source": "SE Ranking MCP (AU)",            // optional one-liner about source
    "rows": [                                     *
      {
        "keyword": "emergency plumber melbourne", *
        "rank": 12,                                  // int, "Unknown", "Not pulled", or null
        "in_three_pack": false,                   *  // true / false / "No 3-pack on SERP" / "Not pulled" / null
        "likely_cause": "GBP completeness gaps + thin on-page content"
      }
    ]
  },
  "action_plan": [                                *
    {
      "priority": "P0",                           *  // P0 / P1 / P2
      "action": "Update Yelp AU phone to +61 3 9876 5432",  *
      "surface": "Yelp Australia",                *
      "impact": "Restores NAP consistency, improves local pack trust signal"  *
    }
  ]
}

The script is forgiving: missing optional fields are handled, and the
document still renders if some sections are sparse.
"""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx is not installed. Install with:\n"
        "    pip install python-docx --break-system-packages\n"
    )
    sys.exit(2)


# --- StudioHawk-ish accent palette (matches the Topic Cluster skill)
HAWK_ORANGE = RGBColor(0xF1, 0x65, 0x1C)
HAWK_DARK = RGBColor(0x1A, 0x1A, 0x1A)
HAWK_GREY = RGBColor(0x55, 0x55, 0x55)
HAWK_RED = RGBColor(0xC0, 0x2B, 0x2B)
HAWK_GREEN = RGBColor(0x1F, 0x7A, 0x3E)
HAWK_AMBER = RGBColor(0xB0, 0x72, 0x1F)


# --- Generic helpers ---------------------------------------------------------

def _shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = HAWK_DARK
        if level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(15)
        else:
            run.font.size = Pt(13)


def _para(doc, text: str, *, bold: bool = False, italic: bool = False, color=None, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _label_value(doc, label: str, value: str) -> None:
    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    value_run = p.add_run(value)
    value_run.font.size = Pt(11)


def _set_header_cell(cell, text: str) -> None:
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _shade_cell(cell, "F1651C")


def _status_colour(status: str) -> RGBColor:
    s = (status or "").strip().lower()
    if s in ("match", "pass", "strong", "healthy", "acceptable"):
        return HAWK_GREEN
    if s in ("partial", "good with gaps", "soft", "amber"):
        return HAWK_AMBER
    if s in ("fail", "failing", "inconsistent", "missing", "poor", "stalled"):
        return HAWK_RED
    return HAWK_DARK


def _priority_fill(priority: str) -> str:
    p = (priority or "").upper().strip()
    if p == "P0":
        return "F4CFCF"  # soft red
    if p == "P1":
        return "F8E4C3"  # soft amber
    if p == "P2":
        return "D7E9D4"  # soft green
    return "EEEEEE"


def _set_cell_text(cell, text: str, *, bold: bool = False, color=None, size: int = 10) -> None:
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = bold
            run.font.size = Pt(size)
            if color is not None:
                run.font.color.rgb = color


# --- Section builders --------------------------------------------------------

def _add_cover(doc, data: dict) -> None:
    title = doc.add_paragraph()
    title_run = title.add_run("Local SEO Audit")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = HAWK_ORANGE

    sub = doc.add_paragraph()
    sub_run = sub.add_run(data["business_name"])
    sub_run.bold = True
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = HAWK_DARK

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"{data['locality']}  ·  {data.get('locale', 'AU')}  ·  "
        f"{data.get('date', date.today().isoformat())}"
    )
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = HAWK_GREY
    doc.add_paragraph()


def _compute_scorecard_totals(rows):
    """Return (earned, total, pct) computed from Pass/Partial/Fail statuses."""
    earned = 0
    total = 0
    for row in rows or []:
        weight = int(row.get("weight", 0) or 0)
        total += weight
        status = (row.get("status") or "").strip()
        if status == "Pass":
            earned += weight
        elif status == "Partial":
            earned += weight / 2  # half credit
    pct = (earned / total * 100) if total else 0
    return earned, total, pct


def _bucket_for_pct(pct: float) -> str:
    if pct >= 85:
        return "Strong"
    if pct >= 70:
        return "Good with gaps"
    if pct >= 50:
        return "Material gaps"
    return "Failing"


def _add_exec_summary(doc, data: dict) -> None:
    es = data.get("executive_summary") or {}
    _heading(doc, "Executive Summary", level=1)

    # Compute the scorecard totals once, here — and reuse in the scorecard
    # footer so the two surfaces can never disagree. JSON values are ignored
    # if a scorecard is present.
    rows = data.get("gbp_scorecard") or []
    if rows:
        earned, total, pct = _compute_scorecard_totals(rows)
        bucket = es.get("gbp_bucket") or _bucket_for_pct(pct)
        # Cache for the scorecard section to reuse
        data["_computed_scorecard"] = {"earned": earned, "total": total, "pct": pct}
        score_str = f"{earned:g}/{total} ({pct:.0f}% — {bucket})"
    else:
        score_str = (
            f"{es.get('gbp_score_points', 'TBC')}/25 "
            f"({es.get('gbp_score_pct', 'TBC')}% — {es.get('gbp_bucket', 'TBC')})"
        )

    _label_value(doc, "GBP completeness", score_str)
    _label_value(doc, "NAP consistency", es.get("nap_summary", "TBC"))
    _label_value(doc, "Reviews", es.get("review_headline", "TBC"))
    _label_value(doc, "3-pack presence", es.get("three_pack_rate", "TBC"))

    actions = es.get("top_actions") or []
    if actions:
        doc.add_paragraph()
        _para(doc, "Top 3 P0 actions", bold=True)
        for a in actions[:3]:
            doc.add_paragraph(a, style="List Bullet")
    doc.add_paragraph()


def _add_scorecard(doc, data: dict) -> None:
    rows = data.get("gbp_scorecard") or []
    _heading(doc, "1. GBP Completeness Scorecard", level=1)
    if not rows:
        _para(doc, "No scorecard data provided.", italic=True, color=HAWK_GREY)
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, label in enumerate(["Field", "Weight", "Status", "Note"]):
        _set_header_cell(table.rows[0].cells[i], label)

    for row in rows:
        weight = int(row.get("weight", 0) or 0)
        status = row.get("status", "Not provided")

        cells = table.add_row().cells
        _set_cell_text(cells[0], row.get("field", ""))
        _set_cell_text(cells[1], str(weight))
        _set_cell_text(cells[2], status, bold=True, color=_status_colour(status))
        _set_cell_text(cells[3], row.get("note", "") or "")

    doc.add_paragraph()
    computed = data.get("_computed_scorecard")
    if computed and computed["total"]:
        _para(
            doc,
            f"Total: {computed['earned']:g} / {computed['total']} points "
            f"({computed['pct']:.0f}% completeness)",
            bold=True,
        )
    else:
        _para(doc, "Total: n/a", italic=True, color=HAWK_GREY)
    doc.add_paragraph()


def _add_nap_report(doc, data: dict) -> None:
    rows = data.get("nap_report") or []
    _heading(doc, "2. NAP Consistency Report", level=1)
    _para(
        doc,
        "Name, Address, and Phone — checked against the canonical listing. "
        "Inconsistencies erode local pack trust signals; fix them in priority order.",
        color=HAWK_GREY,
    )

    if not rows:
        _para(doc, "No citation data provided.", italic=True, color=HAWK_GREY)
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    for i, label in enumerate(["Directory", "Status", "Detail"]):
        _set_header_cell(table.rows[0].cells[i], label)

    # Sort: Inconsistent first, then Missing, then Match
    order = {"Inconsistent": 0, "Missing": 1, "Match": 2}
    rows_sorted = sorted(rows, key=lambda r: order.get(r.get("status", "Match"), 3))

    for r in rows_sorted:
        cells = table.add_row().cells
        _set_cell_text(cells[0], r.get("directory", ""), bold=True)
        _set_cell_text(cells[1], r.get("status", ""), bold=True, color=_status_colour(r.get("status", "")))

        # Build the detail cell content
        detail_lines = []
        if r.get("status") == "Inconsistent":
            if r.get("name_on_listing"):
                detail_lines.append(f"Name on listing: {r['name_on_listing']}")
            if r.get("address_on_listing"):
                detail_lines.append(f"Address on listing: {r['address_on_listing']}")
            if r.get("phone_on_listing"):
                detail_lines.append(f"Phone on listing: {r['phone_on_listing']}")
            if r.get("discrepancy"):
                detail_lines.append(f"Discrepancy: {r['discrepancy']}")
            if r.get("fix_instruction"):
                detail_lines.append(f"Fix: {r['fix_instruction']}")
        elif r.get("status") == "Missing":
            detail_lines.append("Business not currently listed.")
            if r.get("fix_instruction"):
                detail_lines.append(f"Fix: {r['fix_instruction']}")
            else:
                detail_lines.append("Fix: Claim listing and submit canonical NAP.")
        else:  # Match
            detail_lines.append("NAP matches canonical listing — no action.")

        cells[2].text = ""
        for i, line in enumerate(detail_lines):
            p = cells[2].paragraphs[0] if i == 0 else cells[2].add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10)
            if line.startswith("Fix:"):
                run.bold = True

    doc.add_paragraph()


def _add_review_analysis(doc, data: dict) -> None:
    ra = data.get("review_analysis") or {}
    _heading(doc, "3. Review Analysis", level=1)

    _label_value(
        doc,
        "Headline",
        f"{ra.get('avg_rating', 'TBC')} from {ra.get('total_reviews', 'TBC')} reviews",
    )

    vel = ra.get("velocity_per_month")
    vel_str = f"{vel:.1f}/month" if isinstance(vel, (int, float)) else "TBC"
    vel_grade = ra.get("velocity_grade", "TBC")
    _label_value(doc, "Review velocity", f"{vel_str} — {vel_grade}")

    rr = ra.get("response_rate_pct")
    rr_str = f"{rr}%" if isinstance(rr, (int, float)) else "TBC"
    _label_value(
        doc, "Owner response rate", f"{rr_str} — {ra.get('response_grade', 'TBC')}"
    )

    if ra.get("sentiment_note"):
        _label_value(doc, "Recent sentiment", ra["sentiment_note"])

    if ra.get("commentary"):
        doc.add_paragraph()
        _para(doc, ra["commentary"])

    # Competitor comparison — only render if data provided
    competitors = ra.get("competitors") or []
    if competitors:
        doc.add_paragraph()
        _para(doc, "Competitor comparison", bold=True)
        ctable = doc.add_table(rows=1, cols=4)
        ctable.style = "Light Grid Accent 1"
        for i, label in enumerate(["Business", "Domain", "Reviews", "Avg rating"]):
            _set_header_cell(ctable.rows[0].cells[i], label)

        # First row is the audited business
        own_row = ctable.add_row().cells
        _set_cell_text(own_row[0], data.get("business_name", "(this business)"), bold=True)
        _set_cell_text(own_row[1], "—")
        _set_cell_text(own_row[2], str(ra.get("total_reviews", "TBC")))
        _set_cell_text(own_row[3], str(ra.get("avg_rating", "TBC")))

        for comp in competitors:
            cells = ctable.add_row().cells
            _set_cell_text(cells[0], comp.get("name", ""))
            _set_cell_text(cells[1], comp.get("domain", ""))
            _set_cell_text(cells[2], str(comp.get("reviews", "TBC")))
            _set_cell_text(cells[3], str(comp.get("avg_rating", "TBC")))

    doc.add_paragraph()


def _add_rankings(doc, data: dict) -> None:
    r = data.get("rankings") or {}
    rows = r.get("rows") or []
    _heading(doc, "4. Local Keyword Ranking + 3-Pack Report", level=1)
    _label_value(doc, "3-pack presence rate", f"{r.get('three_pack_rate_pct', 'TBC')}%")

    source = r.get("data_source")
    if source:
        _para(doc, f"Data source: {source}", italic=True, color=HAWK_GREY, size=10)

    if not rows:
        _para(doc, "No target keywords provided.", italic=True, color=HAWK_GREY)
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, label in enumerate(["Keyword", "Rank", "3-Pack?", "Likely cause / note"]):
        _set_header_cell(table.rows[0].cells[i], label)

    for kw in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], kw.get("keyword", ""))
        rank = kw.get("rank")
        if rank in (None, "", "Unknown"):
            rank_str = "Unknown"
        elif rank in ("Not pulled", "TBC"):
            rank_str = str(rank)
        else:
            rank_str = str(rank)
        _set_cell_text(cells[1], rank_str)
        in_three = kw.get("in_three_pack")
        if in_three is True:
            _set_cell_text(cells[2], "Yes", bold=True, color=HAWK_GREEN)
        elif in_three is False:
            _set_cell_text(cells[2], "No", bold=True, color=HAWK_RED)
        elif in_three == "No 3-pack on SERP":
            _set_cell_text(cells[2], "No 3-pack on SERP", color=HAWK_GREY)
        elif in_three == "Not pulled":
            _set_cell_text(cells[2], "Not pulled", color=HAWK_AMBER)
        else:
            _set_cell_text(cells[2], "Unknown", color=HAWK_GREY)
        _set_cell_text(cells[3], kw.get("likely_cause", "") or "")

    doc.add_paragraph()


def _add_action_plan(doc, data: dict) -> None:
    actions = data.get("action_plan") or []
    _heading(doc, "5. Prioritised Action Plan", level=1)
    _para(
        doc,
        "Ordered by impact within each priority bucket. "
        "P0 = fix this week, P1 = fix this month, P2 = fix this quarter.",
        color=HAWK_GREY,
    )

    if not actions:
        _para(doc, "No actions provided.", italic=True, color=HAWK_GREY)
        return

    # Sort by P0 -> P1 -> P2 -> other
    order = {"P0": 0, "P1": 1, "P2": 2}
    actions_sorted = sorted(actions, key=lambda a: order.get((a.get("priority") or "").upper(), 99))

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, label in enumerate(["Priority", "Action", "Surface", "Expected impact"]):
        _set_header_cell(table.rows[0].cells[i], label)

    for a in actions_sorted:
        cells = table.add_row().cells
        priority = (a.get("priority") or "").upper()
        _set_cell_text(cells[0], priority, bold=True)
        _shade_cell(cells[0], _priority_fill(priority))
        _set_cell_text(cells[1], a.get("action", ""))
        _set_cell_text(cells[2], a.get("surface", ""))
        _set_cell_text(cells[3], a.get("impact", ""))

    doc.add_paragraph()


def _add_footer(doc) -> None:
    doc.add_paragraph()
    footer = doc.add_paragraph()
    run = footer.add_run(
        "Generated by the Hawk Academy Local SEO Checker. "
        "Audit reflects information provided by the auditor at the time of the run — "
        "verify with the live listings before sending to a client."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = HAWK_GREY


# --- Main --------------------------------------------------------------------

def build_doc(data: dict, output_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_cover(doc, data)
    _add_exec_summary(doc, data)
    _add_scorecard(doc, data)
    _add_nap_report(doc, data)
    _add_review_analysis(doc, data)
    _add_rankings(doc, data)
    _add_action_plan(doc, data)
    _add_footer(doc)

    doc.save(str(output_path))


def main() -> int:
    if len(sys.argv) != 3:
        sys.stderr.write("Usage: python3 build_local_seo_docx.py audit.json output.docx\n")
        return 1
    audit_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    if not audit_path.exists():
        sys.stderr.write(f"ERROR: input file not found: {audit_path}\n")
        return 1
    try:
        data = json.loads(audit_path.read_text())
    except json.JSONDecodeError as e:
        sys.stderr.write(f"ERROR: invalid JSON in {audit_path}: {e}\n")
        return 1
    build_doc(data, output_path)
    print(f"Wrote {output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
