#!/usr/bin/env python3
"""
build_cluster_docx.py

Builds a Hawk Academy topic-cluster Word document from a JSON definition.

Usage:
    python3 build_cluster_docx.py cluster.json output.docx

JSON schema (keys marked * are required):

{
  "pillar_topic": "Email Marketing",          *
  "domain": "example.com.au",                 *
  "locale": "AU",                                // optional, defaults to "AU"
  "date": "2026-05-13",                          // optional, defaults to today
  "blog_section": {                           *
    "status": "exists" | "missing",           *
    "pattern": "/blog/"                       *  // existing pattern, or recommended pattern if missing
  },
  "pillar": {                                 *
    "title": "The Complete Guide to Email Marketing",  *
    "target_keyword": "email marketing",      *
    "msv": 18000,                                // null if TBC
    "cpc": 4.25,                                  // null if TBC
    "intent": "informational",                    // "TBC" if no data
    "kd": 65,                                     // keyword difficulty 0-100, null if TBC
    "recommended_url": "/guides/email-marketing/",  *  // existing URL or proposed slug
    "is_existing": false,                       *  // true = refresh, false = new
    "word_count_target": 2500,                    // optional
    "h2_sections": ["...","..."],             *
    "notes": "Optional copywriter notes"
  },
  "spokes": [                                 *  // 8-10 entries
    {
      "title": "Email Subject Line Best Practices",  *
      "target_keyword": "email subject lines",       *
      "msv": 1300,                                  // null if TBC
      "cpc": 2.10,                                  // null if TBC
      "intent": "informational",                *  // or "TBC"
      "kd": 45,                                     // null if TBC
      "recommended_url": "/blog/email-subject-lines/",  *
      "is_existing": false,                     *  // true = refresh existing, false = new
      "priority": 1                             *  // 1 = highest
    }
  ],
  "cross_spoke_links": [                         // optional
    ["Spoke A title", "Spoke B title"]
  ],
  "production_schedule": [                    *
    {"phase": 0, "task": "Build /blog/ section", "type": "Setup"},
    {"phase": 1, "task": "Build pillar page: ...", "type": "Pillar"},
    {"phase": 2, "task": "Write spoke: ...", "type": "Spoke-new"}
  ]
}

The script is intentionally forgiving: missing optional fields are handled, and
the document still renders if some sections are sparse.
"""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx is not installed. Install with:\n"
        "    pip install python-docx --break-system-packages\n"
    )
    sys.exit(2)


# --- StudioHawk-ish accent palette
HAWK_ORANGE = RGBColor(0xF1, 0x65, 0x1C)
HAWK_DARK = RGBColor(0x1A, 0x1A, 0x1A)
HAWK_GREY = RGBColor(0x55, 0x55, 0x55)


def _shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = HAWK_DARK
        if level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(15)
        else:
            run.font.size = Pt(13)


def _para(doc: Document, text: str, *, bold: bool = False, italic: bool = False, color=None, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    value_run = p.add_run(value)
    value_run.font.size = Pt(11)


def _fmt_number(value, prefix: str = "", suffix: str = "") -> str:
    if value is None or value == "":
        return "TBC"
    try:
        n = float(value)
        if n == int(n):
            return f"{prefix}{int(n):,}{suffix}"
        return f"{prefix}{n:,.2f}{suffix}"
    except (ValueError, TypeError):
        return str(value)


def _fmt_intent(value) -> str:
    if not value or value == "TBC":
        return "TBC"
    return str(value).strip().title()


def _set_header_cell(cell, text: str) -> None:
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _shade_cell(cell, "F1651C")


def build_doc(data: dict, output_path: Path) -> None:
    pillar = data["pillar"]
    spokes = data["spokes"]
    cross_links = data.get("cross_spoke_links", []) or []
    schedule = data.get("production_schedule", []) or []
    blog = data.get("blog_section", {"status": "exists", "pattern": "/blog/"})

    doc = Document()

    # base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Cover block ---
    title = doc.add_paragraph()
    title_run = title.add_run("Topic Cluster Plan")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = HAWK_ORANGE

    sub = doc.add_paragraph()
    sub_run = sub.add_run(data["pillar_topic"])
    sub_run.bold = True
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = HAWK_DARK

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"{data['domain']}  ·  {data.get('locale', 'AU')}  ·  "
        f"{data.get('date', date.today().isoformat())}"
    )
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = HAWK_GREY

    # Blog section status callout
    if blog.get("status") == "missing":
        callout = doc.add_paragraph()
        callout_run = callout.add_run(
            f"⚠ No content section detected on this site. "
            f"Build a {blog.get('pattern', '/blog/')} section before publishing spokes — see Phase 0 in the Production Schedule."
        )
        callout_run.bold = True
        callout_run.font.size = Pt(10)
        callout_run.font.color.rgb = HAWK_ORANGE

    doc.add_paragraph()  # spacer

    # --- 1. Pillar Page Outline ---
    _heading(doc, "1. Pillar Page Outline", level=1)

    _label_value(doc, "Pillar title", pillar["title"])
    _label_value(doc, "Target keyword", pillar["target_keyword"])
    _label_value(doc, "MSV", _fmt_number(pillar.get("msv")))
    _label_value(doc, "CPC", _fmt_number(pillar.get("cpc"), prefix="$"))
    _label_value(doc, "Intent", _fmt_intent(pillar.get("intent")))
    if pillar.get("kd") is not None:
        _label_value(doc, "Keyword difficulty", _fmt_number(pillar.get("kd"), suffix="/100"))
    _label_value(doc, "Recommended URL", pillar.get("recommended_url", "(TBC)"))
    _label_value(doc, "Action", "Refresh / optimise existing pillar" if pillar.get("is_existing") else "Build new pillar page")
    if pillar.get("word_count_target"):
        _label_value(doc, "Word count target", f"{pillar['word_count_target']:,}+ words")
    if pillar.get("notes"):
        _para(doc, pillar["notes"], italic=True, color=HAWK_GREY)

    _heading(doc, "Suggested H2 sections", level=2)
    for h2 in pillar["h2_sections"]:
        p = doc.add_paragraph(h2, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_paragraph()

    # --- 2. Spoke Topics ---
    _heading(doc, "2. Spoke Topics", level=1)
    _para(
        doc,
        f"{len(spokes)} spokes mapped against the cluster. "
        f"Each spoke targets one specific long-tail keyword and links back to the pillar.",
        color=HAWK_GREY,
    )

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = ["#", "Spoke title", "Target keyword", "MSV", "Intent", "Recommended URL"]
    for i, label in enumerate(headers):
        _set_header_cell(table.rows[0].cells[i], label)

    sorted_spokes = sorted(spokes, key=lambda s: s.get("priority", 99))
    for idx, spoke in enumerate(sorted_spokes, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = spoke["title"]
        row[2].text = spoke["target_keyword"]
        row[3].text = _fmt_number(spoke.get("msv"))
        row[4].text = _fmt_intent(spoke.get("intent"))
        row[5].text = spoke.get("recommended_url", "TBC")
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    # CPC + KD appendix below the table (kept light — not the main table)
    cpc_kd_present = any(s.get("cpc") is not None or s.get("kd") is not None for s in spokes)
    if cpc_kd_present:
        doc.add_paragraph()
        _heading(doc, "Spoke keyword details (CPC + difficulty)", level=2)
        appendix = doc.add_table(rows=1, cols=4)
        appendix.style = "Light Grid Accent 1"
        for i, label in enumerate(["#", "Target keyword", "CPC", "Difficulty"]):
            _set_header_cell(appendix.rows[0].cells[i], label)
        for idx, spoke in enumerate(sorted_spokes, start=1):
            row = appendix.add_row().cells
            row[0].text = str(idx)
            row[1].text = spoke["target_keyword"]
            row[2].text = _fmt_number(spoke.get("cpc"), prefix="$")
            row[3].text = _fmt_number(spoke.get("kd"), suffix="/100") if spoke.get("kd") is not None else "TBC"
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

    doc.add_paragraph()

    # --- 3. Internal Linking Map ---
    _heading(doc, "3. Internal Linking Map", level=1)
    _para(doc, "Pillar ⇄ spokes", bold=True)
    p = doc.add_paragraph(
        f"The pillar page links out to all {len(spokes)} spokes. "
        f"Every spoke links back to the pillar with descriptive anchor text "
        f"(e.g. \"the complete guide to {pillar['target_keyword']}\")."
    )
    for run in p.runs:
        run.font.size = Pt(11)

    if cross_links:
        doc.add_paragraph()
        _para(doc, "Cross-spoke links", bold=True)
        for a, b in cross_links:
            doc.add_paragraph(f"{a}  ⇄  {b}", style="List Bullet")
    else:
        _para(
            doc,
            "No cross-spoke links recommended for v1 — keep the structure clean. "
            "Add organically as you write each spoke.",
            italic=True,
            color=HAWK_GREY,
        )

    doc.add_paragraph()

    # --- 4. Production Schedule ---
    _heading(doc, "4. Production Schedule", level=1)
    if not schedule:
        _para(doc, "Schedule not specified.", italic=True, color=HAWK_GREY)
    else:
        sched_table = doc.add_table(rows=1, cols=3)
        sched_table.style = "Light Grid Accent 1"
        for i, label in enumerate(["Phase", "Task", "Type"]):
            _set_header_cell(sched_table.rows[0].cells[i], label)

        for item in schedule:
            row = sched_table.add_row().cells
            phase = item.get("phase", item.get("week", "—"))
            row[0].text = str(phase)
            row[1].text = item.get("task", "")
            row[2].text = item.get("type", "")
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

    # --- Footer note ---
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Generated by the Hawk Academy Topic Cluster Builder. "
        "Keyword data sourced via SE Ranking MCP — verify the locale before client delivery."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = HAWK_GREY

    doc.save(str(output_path))


def main() -> int:
    if len(sys.argv) != 3:
        sys.stderr.write("Usage: python3 build_cluster_docx.py cluster.json output.docx\n")
        return 1
    cluster_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    if not cluster_path.exists():
        sys.stderr.write(f"ERROR: input file not found: {cluster_path}\n")
        return 1
    try:
        data = json.loads(cluster_path.read_text())
    except json.JSONDecodeError as e:
        sys.stderr.write(f"ERROR: invalid JSON in {cluster_path}: {e}\n")
        return 1
    build_doc(data, output_path)
    print(f"Wrote {output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
