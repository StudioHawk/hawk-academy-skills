#!/usr/bin/env python3
"""
build_ux_report.py — Render a UX Diagnosis Report (.docx) from a findings JSON.

Usage:
  python3 build_ux_report.py findings.json /path/to/output.docx

Deps: python-docx  (pip install python-docx --break-system-packages)

JSON schema (only "client" required; the script renders whatever it's given):
{
  "client": "Fit For Life Health & Fitness",
  "url": "https://www.fitforlifefitness.com.au/",
  "date": "2026-06-05",
  "data_sources": ["Microsoft Clarity", "GA4"],         // what the diagnosis is based on
  "summary": "3-5 sentence overall read of the UX health...",
  "quick_wins": ["Fix the dead-click on the hero image", "..."],
  "findings": [
    {
      "area": "Homepage hero",
      "signal": "Dead clicks on the hero image (Clarity, 412 sessions)",
      "evidence": "18% of sessions clicked the hero expecting a link",
      "likely_cause": "Image looks clickable but isn't linked",
      "recommendation": "Link the hero to /personal-training or remove the affordance",
      "severity": "High",            // High / Medium / Low
      "effort": "Low",               // Low / Medium / High
      "impact": "High",              // Low / Medium / High
      "confidence": "Confirmed"      // Confirmed / Hypothesis
    }
  ]
}
"""
import json, sys, datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    sys.stderr.write("Install python-docx: pip install python-docx --break-system-packages\n"); sys.exit(2)

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLUE = RGBColor(0x25, 0x63, 0xEB)
RED = RGBColor(0xDC, 0x26, 0x26)
AMBER = RGBColor(0xD9, 0x77, 0x06)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
GREY = RGBColor(0x64, 0x74, 0x8B)

SEV_COLOR = {"High": "DC2626", "Medium": "D97706", "Low": "16A34A"}


def sev_rank(f):
    s = {"High": 0, "Medium": 1, "Low": 2}.get(f.get("severity", "Medium"), 1)
    # within severity, surface high-impact / low-effort first
    imp = {"High": 0, "Medium": 1, "Low": 2}.get(f.get("impact", "Medium"), 1)
    eff = {"Low": 0, "Medium": 1, "High": 2}.get(f.get("effort", "Medium"), 1)
    return (s, imp, eff)


def run(p, text, bold=False, size=11, color=None, italic=False):
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.name = "Arial"; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r


def H(doc, text, size=16, color=NAVY, before=12, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    run(p, text, bold=True, size=size, color=color); return p


def shade(cell, hexfill):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), hexfill); tcPr.append(sh)


def main():
    if len(sys.argv) < 3:
        sys.stderr.write("Usage: build_ux_report.py findings.json output.docx\n"); sys.exit(1)
    c = json.load(open(sys.argv[1], encoding="utf-8")); out = sys.argv[2]
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)

    # title
    H(doc, "UX Diagnosis Report", 26, NAVY, before=0)
    p = doc.add_paragraph(); run(p, c.get("client", ""), bold=True, size=14, color=BLUE)
    meta = "  ·  ".join([x for x in [c.get("url", ""), c.get("date", ""),
                                     "Sources: " + ", ".join(c.get("data_sources", []))] if x])
    p = doc.add_paragraph(); run(p, meta, size=10, color=GREY)

    # summary
    if c.get("summary"):
        H(doc, "Summary", 15, BLUE)
        p = doc.add_paragraph(); run(p, c["summary"])

    # quick wins
    if c.get("quick_wins"):
        H(doc, "Fix these first (quick wins)", 15, BLUE)
        for q in c["quick_wins"]:
            p = doc.add_paragraph(style="List Bullet"); run(p, q)

    # findings table
    findings = sorted(c.get("findings", []), key=sev_rank)
    if findings:
        H(doc, "Prioritised findings", 15, BLUE)
        cols = ["Area", "Signal & evidence", "Likely cause", "Recommended fix", "Sev", "Effort", "Impact"]
        t = doc.add_table(rows=1, cols=len(cols)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, col in enumerate(cols):
            cell = t.rows[0].cells[i]; shade(cell, "1B2A4A")
            pr = cell.paragraphs[0]; run(pr, col, bold=True, size=9, color=RGBColor(0xFF,0xFF,0xFF))
        for f in findings:
            cells = t.add_row().cells
            sig = f.get("signal", "")
            if f.get("evidence"): sig += " — " + f["evidence"]
            conf = f.get("confidence", "")
            cause = f.get("likely_cause", "")
            if conf: cause = "[%s] %s" % (conf, cause)
            vals = [f.get("area", ""), sig, cause, f.get("recommendation", ""),
                    f.get("severity", ""), f.get("effort", ""), f.get("impact", "")]
            for i, v in enumerate(vals):
                pr = cells[i].paragraphs[0]; run(pr, str(v), size=9)
            # colour the severity cell
            sv = f.get("severity", "Medium")
            shade(cells[4], SEV_COLOR.get(sv, "D97706"))
            cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            cells[4].paragraphs[0].runs[0].bold = True

    # method note
    H(doc, "How to read this", 13, BLUE)
    p = doc.add_paragraph()
    run(p, "Behavioural data shows what users do, not always why. Findings marked [Confirmed] were "
            "verified in session recordings or corroborated across tools; [Hypothesis] findings should be "
            "validated with a recording, on-site survey, or A/B test before a large rebuild.", size=10, color=GREY, italic=True)

    doc.save(out); print("Wrote", out)


if __name__ == "__main__":
    main()
